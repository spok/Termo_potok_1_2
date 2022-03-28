import os
import sys
import matplotlib.pyplot as plt
import main_window
from data_calc import MyData
from PyQt5.QtWidgets import QTableWidgetItem, QFileDialog, QInputDialog, QMessageBox
from PyQt5 import QtGui, QtWidgets
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import matplotlib.dates as mdates
from pyqtgraph.Qt import QtGui
import pyqtgraph as pg
import docx
import numpy as np

class MyWindow(QtWidgets.QMainWindow, QtWidgets.QWidget, main_window.Ui_MainWindow):
    """Основной класс программы"""

    def __init__(self):
        # Инициализация родителей
        super().__init__()
        self.setupUi(self)

        # Устанавливаем размер таблицы
        self.table.setColumnCount(11)  # Устанавливаем 11 колонок, 1 столбец с датой и 10 столбцов с каналами
        self.table.setRowCount(10)  # и десять строк в таблице для начала
        self.table2.setColumnCount(10)  # Устанавливаем 10 столбцов для каналов
        self.table2.setRowCount(3)  # и 3 строки для данных по каналам

        # настройка таблицы для расчета конструкции
        self.table_calc.setColumnCount(6)
        self.table_calc.setRowCount(4)
        self.table_calc.setVerticalHeaderLabels(['q', 'Твн.пов', 'Твн', 'Тнар'])
        self.table_calc.setHorizontalHeaderLabels(['Номер\nканала', 'Среднее\nзначение',
                                                   'Сдвиг на\nединицы', 'Сдвиг на\n десятые',
                                                   'Масштаб', 'Итоговое'])
        w = 65
        self.table_calc.setColumnWidth(0, w)
        self.table_calc.setColumnWidth(1, w)
        self.table_calc.setColumnWidth(2, w)
        self.table_calc.setColumnWidth(3, w)
        self.table_calc.setColumnWidth(4, w)
        self.table_calc.setColumnWidth(5, w)
        self.table_calc.setFixedHeight(170)
        # Добавление элементов управления в таблицу
        for i in range(0, 4):
            self.table_calc.setItem(i, 1, QTableWidgetItem())
            # настройка элементов выпадающего списка с номерами каналов
            elem_combo = QtWidgets.QComboBox()
            elem_combo.activated.connect(self.change_combobox)
            self.table_calc.setCellWidget(i, 0, elem_combo)
            self.table_calc.setItem(i, 1, QTableWidgetItem())
            # настройка спинбокса для шага на единицы
            elem_spin = QtWidgets.QSpinBox()
            elem_spin.setValue(0)
            elem_spin.setSingleStep(1)
            elem_spin.setMinimum(-50)
            elem_spin.setMaximum(50)
            elem_spin.valueChanged.connect(self.change_step1)  # изменение шага по времени
            self.table_calc.setCellWidget(i, 2, elem_spin)
            # настройка спинбокса для шага на десятые единицы
            elem_dspin = QtWidgets.QDoubleSpinBox(decimals=3)
            elem_dspin.setValue(0.000)
            elem_dspin.setSingleStep(0.100)
            elem_dspin.setMinimum(-50)
            elem_dspin.setMaximum(50)
            elem_dspin.valueChanged.connect(self.change_step2)  # изменение шага по времени
            self.table_calc.setCellWidget(i, 3, elem_dspin)
            # настройка спинбокса для изменения масштабного коэфициента
            elem_spin = QtWidgets.QSpinBox()
            elem_spin.setValue(0)
            elem_spin.setSingleStep(1)
            elem_spin.setMinimum(0)
            elem_spin.setMaximum(50)
            elem_spin.valueChanged.connect(self.change_koef)  # изменение шага по времени
            self.table_calc.setCellWidget(i, 4, elem_spin)
            self.table_calc.setItem(i, 5, QTableWidgetItem())

        # настройка таблицы для сопротивлений теплопередаче
        self.table_r.setColumnCount(3)
        self.table_r.setRowCount(3)
        self.table_r.setVerticalHeaderLabels(['Конструкция №1', 'Конструкция №2', 'Конструкция №3'])
        self.table_r.setHorizontalHeaderLabels(
            ['Сопротивление\nтеплопередаче R1', 'Сопротивление\nтеплопередаче R2', 'Погрешность\nизмерений'])
        w = 120
        self.table_r.setColumnWidth(0, w)
        self.table_r.setColumnWidth(1, w)
        self.table_r.setColumnWidth(2, w)
        self.table_r.setFixedHeight(150)

        # Устанавливаем заголовки таблиц
        self.table.setHorizontalHeaderLabels(["Дата", "q1", "q2", "q3", "t1", "t2", "t3", "t4", "t5", "t6", "t7"])
        self.table2.setHorizontalHeaderLabels(["q1", "q2", "q3", "t1", "t2", "t3", "t4", "t5", "t6", "t7"])
        self.table2.setVerticalHeaderLabels(["Среднее", "Минимальное", "Максимальное"])

        # делаем ресайз колонок по содержимому
        self.table.resizeColumnsToContents()
        self.table2.resizeColumnsToContents()

        # создание канвы для графика
        self.graphWidget = pg.PlotWidget()
        self.verticalLayout_5.addWidget(self.graphWidget)
        self.graphWidget.clear()
        self.graphWidget.setBackground('w')
        self.graph_line = []
        self.graph_dot = []
        self.graphWidget2 = pg.PlotWidget()
        self.graphWidget2.clear()
        self.graphWidget2.setBackground('w')
        self.graphWidget3 = pg.PlotWidget()
        self.graphWidget3.clear()
        self.graphWidget3.setBackground('w')
        self.verticalLayout_6.addWidget(self.graphWidget2)
        self.verticalLayout_7.addWidget(self.graphWidget3)

        # добавление элементов в выпадающий список
        for i in range(3):
            self.comboBox.addItem('Конструкция № ' + str(i + 1))
        self.comboBox.setCurrentIndex(0)

        # создание ползунков для выбора диапазона графика и задание стиля оформления
        self.slider_left.setStyleSheet("""
            QSlider{
                background: #E3DEE2;
            }
            QSlider::groove:horizontal {  
                height: 10px;
                margin: 0px;
                border-radius: 5px;
                background: #B0AEB1;
            }
            QSlider::handle:horizontal {
                background: #fff;
                border: 1px solid #E3DEE2;
                width: 17px;
                margin: -5px 0; 
                border-radius: 8px;
            }
            QSlider::add-page:qlineargradient {
                background: #3B99FC;
                border-radius: 5px;
            }
        """)
        self.slider_left.setEnabled(False)
        self.slider_right.setStyleSheet("""
            QSlider{
                background: #E3DEE2;
            }
            QSlider::groove:horizontal {  
                height: 10px;
                margin: 0px;
                border-radius: 5px;
                background: #B0AEB1;
            }
            QSlider::handle:horizontal {
                background: #fff;
                border: 1px solid #E3DEE2;
                width: 17px;
                margin: -5px 0; 
                border-radius: 8px;
            }
            QSlider::sub-page:qlineargradient {
                background: #3B99FC;
                border-radius: 5px;
            }
        """)
        self.slider_right.setEnabled(False)
        self.error_dialog = QtWidgets.QErrorMessage()
        # настройка обработчиков событий
        self.table.blockSignals(True)
        self.button1.clicked.connect(self.showdialog)
        self.button2.clicked.connect(self.showdialog2)
        self.spinBox.valueChanged.connect(self.change_time)  # изменение шага по времени
        self.slider_left.valueChanged.connect(self.show_label_left)
        self.slider_right.valueChanged.connect(self.show_label_right)
        self.table.itemSelectionChanged.connect(self.show_graph)
        self.table.cellChanged.connect(self.cell_edit)
        self.tabWidget.currentChanged.connect(lambda: self.show_constr(0))
        # смена конструкции в окне
        self.comboBox.currentIndexChanged.connect(self.show_constr)
        # смена наименования конструкции
        self.lineEdit.textEdited.connect(self.change_name)
        # сохранение графика в файл
        self.button6.clicked.connect(self.export_base)
        # сохранение и открытие файлов в формате программы
        self.button4.clicked.connect(self.save_base)
        self.button5.clicked.connect(self.load_base)
        self.table.blockSignals(False)

    def showdialog(self):
        """отображение диалога открытия файла и загрузка данных из таблицы Excel"""
        fname = QFileDialog.getOpenFileName(self, 'Open file', '', '*.xlsx')[0]
        self.setWindowTitle("Обработка данных датчиков - " + os.path.basename(fname))
        constr.word_file = os.path.splitext(fname)[0]
        if fname != '':
            constr.load_excel(fname)
            # перерисовка всего окна
            self.redraw_window()
            self.change_resulf()

    def showdialog2(self):
        """отображение диалога открытия файла и загрузка данных из исходного файла потоков"""
        fname = QFileDialog.getOpenFileName(self, 'Open file', '', '*.db')[0]
        self.setWindowTitle("Обработка данных датчиков - " + os.path.basename(fname))
        constr.word_file = os.path.splitext(fname)[0]
        if fname != '':
            series = 1
            text, ok = QInputDialog.getText(self, 'Выбор серии',
                                            'В файле с данными может быть больше одной серии с данными.\n Необходимо выбрать серию: ')
            if ok:
                series = int(text)
            constr.load_potokfile(fname, series)
            # перерисовка всего окна
            self.redraw_window()
            self.change_resulf()

    def redraw_window(self):
        """Перерисовка всех элементов формы с новыми данными"""
        self.slider_left.setEnabled(True)
        self.slider_right.setEnabled(True)
        self.set_minmax_labels()
        self.show_label()
        self.show_data()
        self.show_data2()
        self.spinBox.setValue(constr.time_step)
        self.table.setCurrentCell(0, 0)
        self.show_graph()
        # Добавление наименований каналов в выпадающий список
        for i in range(4):
            elem = self.table_calc.cellWidget(i, 0)
            elem.addItems(constr.data_label)

    def set_minmax_labels(self):
        """Настройка диапазона слайдеров"""
        l = constr.get_count_element()
        # блокировка обработки сигналов слайдеров
        self.slider_left.blockSignals(True)
        self.slider_right.blockSignals(True)
        # настройка диапазона слайдеров
        self.slider_left.setMinimum(1)
        self.slider_left.setMaximum(l)
        self.slider_right.setMinimum(1)
        self.slider_right.setMaximum(l)
        # снятие блокировки обработки сигналов слайдеров
        self.slider_left.blockSignals(False)
        self.slider_right.blockSignals(False)

    def show_label(self):
        """Изменение настроек слайдеров"""
        # блокировка обработки сигналов слайдеров
        self.slider_left.blockSignals(True)
        self.slider_right.blockSignals(True)
        # настройка левого слайдера
        border = constr.get_border('left') + 1
        self.label_left.setText('Начало измерений - ' + str(border))
        self.slider_left.setValue(border)
        # настройка правого слайдера
        border = constr.get_border('right')
        self.label_right.setText('Конец измерений - ' + str(border))
        self.slider_right.setValue(border)
        # снятие блокировки обработки сигналов слайдеров
        self.slider_left.blockSignals(False)
        self.slider_right.blockSignals(False)

    def show_label_left(self, value):
        """Установка левой границы диапазона измерений"""
        time_str = constr.date_time[value - 1].strftime('%d.%m.%y %H:%M')
        self.label_left.setText(f'Начало измерений - {str(value)} ({time_str})')
        constr.left_border = value - 1
        if constr.left_border > self.table.currentRow():
            self.table.setCurrentCell(constr.left_border, self.table.currentColumn())
        self.update_graph()
        self.show_data2()
        self.change_resulf()

    def show_label_right(self, value):
        """Установка правой границы диапазона измерений"""
        time_str = constr.date_time[value - 1].strftime('%d.%m.%y %H:%M')
        self.label_right.setText(f'Конец измерений - {str(value)} ({time_str})')
        self.slider_left.setMaximum(value)
        constr.right_border = value
        if constr.right_border < self.table.currentRow():
            self.table.setCurrentCell(constr.right_border, self.table.currentColumn())
        self.update_graph()
        self.show_data2()
        self.change_resulf()

    def paint_cell(self, cur_col, cur_row, cur_data):
        """изменение фона ячейки в зависимости от отклонения значения ячейки от среднего"""
        self.table.setItem(cur_row, cur_col + 1, QTableWidgetItem(str(cur_data)))
        # измениние заливки ячейки в зависимости от отклонения от среднего значения
        s = constr.data_dev[cur_col][cur_row]
        if s > 40.0:
            self.table.item(cur_row, cur_col + 1).setBackground(QtGui.QColor(200, 235, 250))
        if s > 50.0:
            self.table.item(cur_row, cur_col + 1).setBackground(QtGui.QColor(148, 216, 246))
        if s > 60.0:
            self.table.item(cur_row, cur_col + 1).setBackground(QtGui.QColor(96, 197, 241))
        if s > 80.0:
            self.table.item(cur_row, cur_col + 1).setBackground(QtGui.QColor(0, 122, 174))
        if s > 90.0:
            self.table.item(cur_row, cur_col + 1).setBackground(QtGui.QColor(0, 81, 116))

    def show_canal(self, canal: int):
        """Отображение в таблице столбца одного канала"""
        self.table.blockSignals(True)
        d = constr.get_canal_data()
        cur_row = 0
        for cur_data in d[canal]:
            self.paint_cell(canal, cur_row, cur_data)
            cur_row += 1
        self.table.blockSignals(False)

    def show_data(self):
        """Вывод показаний датчиков в основную таблицу окна"""
        # блокировка реагирования на изменение содержания таблицы
        self.table.blockSignals(True)
        # очистка содержания таблицы
        self.table.clear()
        # изменение количества строк в таблице
        self.table.setRowCount(constr.get_count_element())
        # оформление шапки таблиц
        canal_type = constr.get_list_canal()
        self.table.setHorizontalHeaderLabels(["Дата"] + canal_type)
        self.table2.setHorizontalHeaderLabels(canal_type)
        # отображение времени измерений
        self.show_time()
        # добавление показаний датчиков
        for j in range(0, 10):
            self.show_canal(j)
        # отображение таблицы и графика
        self.table.resizeColumnsToContents()
        self.table2.resizeColumnsToContents()
        self.table.blockSignals(False)

    def show_data2(self):
        """Заполнение таблицы для статистических данных"""
        for j in range(0, 10):
            # вывод данных в таблицу
            m = constr.get_mid_canal(j, 'mid')
            self.table2.setItem(0, j, QTableWidgetItem(f'{m:.2f}'))
            m = constr.get_mid_canal(j, 'min')
            self.table2.setItem(1, j, QTableWidgetItem(str(m)))
            m = constr.get_mid_canal(j, 'max')
            self.table2.setItem(2, j, QTableWidgetItem(str(m)))

    def show_time(self):
        """отображение в таблице даты и времени измерений"""
        # добавление времени измерений
        self.table.blockSignals(True)
        for i, cur_time in enumerate(constr.date_time):
            # добавление даты с преобразованием формата
            time_str = cur_time.strftime('%d.%m.%y %H:%M')
            self.table.setItem(i, 0, QTableWidgetItem(time_str))
        self.table.blockSignals(False)

    def change_time(self, value):
        """Изменение времени измерений на заданный шаг"""
        constr.set_date_time(value)
        self.show_time()

    def show_graph(self):
        """Отрисовка графика"""
        if len(constr.data) > 0:
            self.graphWidget.clear()
            self.graphWidget.setBackground('w')
            pen = [pg.mkPen(color=(0, 32, 96), width=2),
                   pg.mkPen(color=(0, 112, 192), width=2),
                   pg.mkPen(color=(0, 176, 240), width=2),
                   pg.mkPen(color=(0, 176, 80)),
                   pg.mkPen(color=(146, 208, 80)),
                   pg.mkPen(color=(170, 174, 235)),
                   pg.mkPen(color=(255, 192, 0)),
                   pg.mkPen(color=(251, 207, 208)),
                   pg.mkPen(color=(128, 128, 128)),
                   pg.mkPen(color=(112, 48, 160))]
            self.graph_line = []
            self.graph_dot = []
            self.graphWidget.addLegend()
            self.graphWidget.showGrid(x=True, y=True)
            self.graphWidget.setYRange(constr.min_value, constr.max_value, padding=0)

            x_cut = list(range(constr.left_border, constr.right_border))
            # определение активной ячейки
            cur_col = self.table.currentColumn()
            cur_row = self.table.currentRow()
            # список с координатами точек активных ячеек
            x_dot = []
            y_dot = []
            # если активная ячейка в столбце дат
            if cur_col < 1:
                for j in range(0, 10):
                    sensor_list = constr.data[j]
                    x_dot.append(cur_row)
                    y_dot.append(sensor_list[cur_row])
                    y = constr.data[j, constr.left_border:constr.right_border]
                    plotname = constr.data_label[j]
                    self.graph_line.append(self.graphWidget.plot(x_cut, y, name=plotname, pen=pen[j]))
                if y_dot:
                    self.graph_dot.append(
                        self.graphWidget.plot(x_dot, y_dot, pen=None, symbol='o', symbolSize=7, symbolBrush=('r')))
            # если активная ячейка в столбце показаний датчика
            else:
                sensor_list = constr.data[cur_col - 1]
                x_dot.append(cur_row)
                y_dot.append(sensor_list[cur_row])
                y = constr.data[cur_col - 1, constr.left_border:constr.right_border]
                plotname = constr.data_label[cur_col - 1]
                self.graph_line.append(self.graphWidget.plot(x_cut, y, name=plotname, pen=pen[cur_col - 1]))
                if y_dot:
                    self.graph_dot.append(
                        self.graphWidget.plot(x_dot, y_dot, pen=None, symbol='o', symbolSize=7, symbolBrush=('r')))

    def show_graph2(self, graph, rez: list, legend: list):
        """Отрисовка графика результатов обработки"""
        if len(rez) > 0:
            graph.clear()
            graph.setBackground('w')
            pen = [pg.mkPen(color=(0, 32, 96), width=2),
                   pg.mkPen(color=(0, 112, 192), width=2),
                   pg.mkPen(color=(0, 176, 240), width=2),
                   pg.mkPen(color=(0, 176, 80), width=2),
                   pg.mkPen(color=(146, 208, 80)),
                   pg.mkPen(color=(170, 174, 235)),
                   pg.mkPen(color=(255, 192, 0)),
                   pg.mkPen(color=(251, 207, 208)),
                   pg.mkPen(color=(128, 128, 128)),
                   pg.mkPen(color=(112, 48, 160))]
            graph.addLegend()
            graph.showGrid(x=True, y=True)
            if len(rez) > 2:
                # поиск вертикальных границ графика
                border_y = constr.get_border_y(rez)
                graph.setYRange(border_y[0], border_y[1], padding=0)
            else:
                graph.setYRange(0, 10, padding=0)
            x_cut = list(range(len(rez[0])))
            # если активная ячейка в столбце дат
            for j in range(len(rez)):
                y = rez[j]
                graph.plot(x_cut, y, name=legend[j], pen=pen[j])

    def update_graph(self):
        """Отрисовка графика"""
        if len(constr.data) > 0:
            cur_col = self.table.currentColumn()
            cur_row = self.table.currentRow()
            x_cut = list(range(constr.left_border, constr.right_border))
            if cur_col < 1:
                for j in range(0, 10):
                    y = constr.data[j, constr.left_border: constr.right_border]
                    self.graph_line[j].setData(x_cut, y)
            else:
                if len(self.graph_line) > 0:
                    y = constr.data[cur_col - 1, constr.left_border:constr.right_border]
                    self.graph_line[0].setData(x_cut, y)

    def cell_edit(self, row: int, col: int):
        """Функция сохранения введенного значения в ячейку в массиве данных"""
        self.table.blockSignals(True)
        if row >= 0 and col > 0:
            cell_text = self.table.item(row, col).text()
            cell_value = float(cell_text)
            constr.edit_cell_data(col - 1, row, cell_value)
            self.paint_cell(col - 1, row, cell_value)
            self.table.setCurrentCell(row, col)
            self.update_graph()
            self.show_data2()
        self.table.blockSignals(False)

    def show_constr(self, index: int):
        """Отображение параметров конструкций в графических элементах вкладки"""
        if constr.get_count_element() > 0 and self.tabWidget.currentIndex() > 0:
            # Отключение обработчика событий
            state = True
            for r in range(4):
                for k in range(6):
                    elem = self.table_calc.cellWidget(r, k)
                    if isinstance(elem, QtWidgets.QComboBox):
                        elem.blockSignals(state)
                    if isinstance(elem, QtWidgets.QSpinBox):
                        elem.blockSignals(state)
                    if isinstance(elem, QtWidgets.QDoubleSpinBox):
                        elem.blockSignals(state)

            # Вывод данных в графические элементы окна
            if index != self.comboBox.currentIndex():
                index = self.comboBox.currentIndex()
            constr.current_canals = index
            cur_canal = constr.canals[index]
            buf = cur_canal.canal_number
            buf2 = cur_canal.canal_value
            for i in range(4):
                elem = self.table_calc.cellWidget(i, 0)
                if isinstance(elem, QtWidgets.QComboBox):
                    elem.setCurrentIndex(buf[i])
                self.table_calc.setItem(i, 1, QTableWidgetItem('{0:.3f}'.format(constr.data_mid[buf[i]])))
                elem = self.table_calc.cellWidget(i, 2)
                if isinstance(elem, QtWidgets.QSpinBox):
                    elem.setValue(cur_canal.canal_koef[i][0])
                elem = self.table_calc.cellWidget(i, 3)
                if isinstance(elem, QtWidgets.QDoubleSpinBox):
                    elem.setValue(cur_canal.canal_koef[i][1])
                elem = self.table_calc.cellWidget(i, 4)
                if isinstance(elem, QtWidgets.QSpinBox):
                    elem.setValue(cur_canal.canal_koef[i][2])
                self.table_calc.setItem(i, 5, QTableWidgetItem('{0:.3f}'.format(buf2[i])))

            self.lineEdit.setText(cur_canal.name)
            for i in range(3):
                text = '{0:.3f}'.format(constr.canals[i].ro1)
                self.table_r.setItem(i, 0, QTableWidgetItem(text))
                text = '{0:.3f}'.format(constr.canals[i].ro2)
                self.table_r.setItem(i, 1, QTableWidgetItem(text))
                text = '{0:.3f}'.format(constr.canals[i].dr)
                self.table_r.setItem(i, 2, QTableWidgetItem(text))

            # Включение обработчика событий
            state = False
            for r in range(4):
                for k in range(6):
                    elem = self.table_calc.cellWidget(r, k)
                    if isinstance(elem, QtWidgets.QComboBox):
                        elem.blockSignals(state)
                    if isinstance(elem, QtWidgets.QSpinBox):
                        elem.blockSignals(state)
                    if isinstance(elem, QtWidgets.QDoubleSpinBox):
                        elem.blockSignals(state)

            # Отрисовка графиков
            self.show_graph2(self.graphWidget2,
                             constr.canals[index].canal,
                             constr.canals[index].canal_name)
            self.show_graph2(self.graphWidget3,
                             constr.canals[index].ro,
                             ['Ro1', 'Ro2'])

    def change_name(self):
        """Обработка изменение названия конструкции, canal - номер конструкции"""
        index = self.comboBox.currentIndex()
        constr.canals[index].name = self.lineEdit.text()

    def change_combobox(self):
        """Обработка изменение номера канала"""
        index = self.comboBox.currentIndex()
        row = self.table_calc.currentRow()
        elem = self.table_calc.cellWidget(row, 0)
        if isinstance(elem, QtWidgets.QComboBox):
            constr.canals[index].canal_number[row] = elem.currentIndex()
            self.table_calc.setItem(row, 1, QTableWidgetItem('{0:.3f}'.format(constr.data_mid[elem.currentIndex()])))
        self.change_resulf()

    def change_step1(self):
        """Обработка изменение шага одного из каналов на единицы"""
        index = self.comboBox.currentIndex()
        row = self.table_calc.currentRow()
        elem = self.table_calc.cellWidget(row, 2)
        if isinstance(elem, QtWidgets.QSpinBox):
            constr.canals[index].canal_koef[row][0] = elem.value()
        self.change_resulf()

    def change_step2(self):
        """Обработка изменение шага одного из каналов на десятые"""
        index = self.comboBox.currentIndex()
        row = self.table_calc.currentRow()
        elem = self.table_calc.cellWidget(row, 3)
        if isinstance(elem, QtWidgets.QDoubleSpinBox):
            constr.canals[index].canal_koef[row][1] = elem.value()
        self.change_resulf()

    def change_koef(self):
        """Обработка изменение масштабного коэффициента"""
        index = self.comboBox.currentIndex()
        row = self.table_calc.currentRow()
        elem = self.table_calc.cellWidget(row, 4)
        if isinstance(elem, QtWidgets.QSpinBox):
            constr.canals[index].canal_koef[row][2] = elem.value()
        self.change_resulf()
        # logger.debug('результат ' +
        # str(constr.canals[index].canal[constr.canals[index].canal_number[row]]))

    def change_resulf(self):
        """Пересчет изменений значений измерений при изменении значений коэффициентов"""
        constr.calc_r()
        # вывод пересчитанных средних значений каналов
        index = self.comboBox.currentIndex()
        self.show_constr(index)

    def save_base(self):
        fname = QFileDialog.getSaveFileName(self, 'Save .json file', '', '*.json')[0]
        result = constr.save_base(fname)
        if result == 'OK':
            QMessageBox.about(self, 'Сохранение файла', 'Файл записан')
        else:
            QMessageBox.about(self, 'Сохранение файла', 'Не удалось записать файл')

    def load_base(self):
        fname = QFileDialog.getOpenFileName(self, 'Load .json file', '', '*.json')[0]
        result = constr.load_base(fname)
        if result == 'OK':
            # перерисовка всего окна
            self.change_time(constr.time_step)
            self.redraw_window()
            self.change_resulf()
        else:
            QMessageBox.about(self, 'Загрузка файла', 'Не удалось считать файл')

    def plot_file(self, rez: list, legend: list, index: int):
        """Сохранение графика потоков в файле"""
        if len(rez)>0:
            #загрузка отметок времени
            x = np.array(constr.date_time[constr.left_border: constr.right_border])
            #настройка отображения графиков
            plt.figure(figsize=(8, 6))
            plt.xlabel('Время проведения измерений')
            if len(rez) > 2:
                plt.ylabel('Температура, °С, тепловой поток Вт/м²')
                filename = 'grafic_potoc' + str(index + 1) + '.jpg'
            else:
                plt.ylabel('Сопротивление теплопередаче, м²°С/Вт')
                filename = 'grafic_ro' + str(index+1) + '.jpg'
            plt.grid(True)
            #настройка горизонтальной оси
            plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%d.%m.%y %H:%M'))
            plt.xticks(rotation=45)
            plt.tick_params(axis='x', which='major', labelsize=8)
            plt.subplots_adjust(left=0.1, bottom=0.165, right=0.97, top=0.995, wspace=0.2, hspace=0.2)

            #загрузка данных датчиков
            if len(rez) > 2:
                colors = ['purple', 'green', 'red', 'blue']
            else:
                colors = ['maroon', 'blue']

            for i, y in enumerate(rez):
                plt.plot(x, y, color=colors[i], linewidth=2.0, label=legend[i])

            plt.legend()
            plt.savefig(filename, dpi=300)
        return filename

    def export_base(self):
        """Вывод результата в вордовский файл"""
        """Экспорт графиков в формате word"""
        doc = docx.Document()
        num_ris = 1
        file_ris = []
        # Генерация рисунков и вставка в вордовский файл
        for i, can in enumerate(constr.canals):
            # График с изменениями значений каналов
            name = self.plot_file(can.canal, can.canal_name, i)
            paragraph = doc.add_picture(name, width=docx.shared.Cm(13))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1
            paragraph.alignment = 1
            if can.name == '':
                buf_text = 'Рис. ' + str(num_ris) + '. Результаты измерений ограждающей конструкции №' + str(i+1)
            else:
                buf_text = 'Рис. ' + str(num_ris) + '. Результаты измерений ограждающей конструкции №' + str(i+1) \
                           + ' (' + can.name + ')'
            paragraph = doc.add_paragraph(buf_text)
            paragraph.alignment = 1
            num_ris += 1
            file_ris.append(name)

            # График с сопротивлениями тепловпередаче
            name = self.plot_file(can.ro, ['Ro1', 'Ro2'], i)
            paragraph = doc.add_picture(name, width=docx.shared.Cm(13))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1
            paragraph.alignment = 1
            if can.name == '':
                buf_text = 'Рис. ' + str(num_ris) + '. Измеренное сопротивление теплопередаче конструкции №' + str(i+1)
            else:
                buf_text = 'Рис. ' + str(num_ris) + '. Измеренное сопротивление теплопередаче конструкции №' \
                           + str(i+1) + ' (' + can.name + ')'
            paragraph = doc.add_paragraph(buf_text)
            paragraph.alignment = 1
            file_ris.append(name)
            num_ris += 1

        # Заголовок таблицы
        paragraph = doc.add_paragraph('')
        paragraph = doc.add_paragraph('Таблица 1. Измеренные параметры ОКЗ')
        paragraph.alignment = 1

        # Вставка таблицы с средними значениями
        table_rows = 4
        table_cols = 5
        table = doc.add_table(rows=table_rows, cols=table_cols)
        table.style = 'Table Grid'
        table.alignment = 1
        cell = table.cell(0, 0)
        cell.text = 'Наименование и расположение исследуемых конструкций'
        cell.paragraphs[0].alignment = 1
        cell = table.cell(0, 1)
        cell.text = 'Средняя температура внутреннего воздуха, °С'
        cell.paragraphs[0].alignment = 1
        cell = table.cell(0, 2)
        cell.text = 'Средняя температура наружного воздуха, °С'
        cell.paragraphs[0].alignment = 1
        cell = table.cell(0, 3)
        cell.text = 'Средняя температура внутренней поверхности, °С'
        cell.paragraphs[0].alignment = 1
        cell = table.cell(0, 4)
        cell.text = 'Средний тепловой поток, Вт/м2'
        cell.paragraphs[0].alignment = 1
        table_rows = 3
        for row, con in enumerate(constr.canals):
            for col in range(table_cols):
                # получаем ячейку таблицы
                cell = table.cell(row+1, col)
                # записываем в ячейку данные
                if col == 0:
                    if con.name == '':
                        buf_text = 'Конструкция №' + str(row + 1)
                    else:
                        buf_text = 'Конструкция №' + str(row + 1) + ' (' + con.name + ')'
                    cell.text = buf_text
                if col == 1:
                    buf = '{0:.3f}'.format(con.canal_value[2])
                    cell.text = buf
                if col == 2:
                    buf = '{0:.3f}'.format(con.canal_value[3])
                    cell.text = buf
                if col == 3:
                    buf = '{0:.3f}'.format(con.canal_value[1])
                    cell.text = buf
                if col == 4:
                    buf = '{0:.3f}'.format(con.canal_value[0])
                    cell.text = buf
                cell.paragraphs[0].alignment = 1

        # Вставка заголовка таблицы
        paragraph = doc.add_paragraph('')
        paragraph = doc.add_paragraph('Таблица 2. Усредненные величины сопротивлений теплопередаче ОКЗ, м2·°С/Вт')
        paragraph.alignment = 1
        # Вставка таблицы с сопротивлением теплопередаче
        table_rows = 4
        table_cols = 4
        table = doc.add_table(rows=table_rows, cols=table_cols)
        table.style = 'Table Grid'
        table.alignment = 1
        cell = table.cell(0, 0)
        cell.text = 'Наименование и расположение исследуемых конструкций'
        cell.paragraphs[0].alignment = 1
        cell = table.cell(0, 1)
        cell.text = 'Нормативное значение'
        cell.paragraphs[0].alignment = 1
        cell = table.cell(0, 2)
        cell.text = 'Проектное значение'
        cell.paragraphs[0].alignment = 1
        cell = table.cell(0, 3)
        cell.text = 'Измеренное значение'
        cell.paragraphs[0].alignment = 1
        table_rows = 3
        for row, con in enumerate(constr.canals):
            for col in range(table_cols):
                # получаем ячейку таблицы
                cell = table.cell(row+1, col)
                # записываем в ячейку данные
                if col == 0:
                    if con.name == '':
                        buf_text = 'Конструкция №' + str(row + 1)
                    else:
                        buf_text = 'Конструкция №' + str(row + 1) + ' (' + con.name + ')'
                    cell.text = buf_text

                if col == 3:
                    buf = '{0:.3f}'.format(con.ro1)
                    buf = buf + '±' + '{0:.3f}'.format(con.dr)
                    cell.text = buf
                cell.paragraphs[0].alignment = 1

        # Сохранение сгенерированного doc файла
        if constr.word_file != '':
            doc.save(constr.word_file + ' - обработанный.docx')
        else:
            constr.word_file = 'export_data.docx'
            doc.save(constr.word_file)

        # Удаление временных файлов
        for f in file_ris:
            try:
                os.remove(f)
            except:
                QMessageBox.about(self, 'Ошибка', 'Невозможно удалить временные файлы')
        QMessageBox.about(self, 'Экспорт', 'Генерация файла завершена')

class Canvas(FigureCanvas):
    def __init__(self, parent=None, width=5, height=5, dpi=100):
        self.fig = Figure(figsize=(width, height), dpi=dpi)
        self.axes = self.fig.add_subplot(111)
        self.plot_ref = []
        self.plot_dot = None
        super(Canvas, self).__init__(self.fig)


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    window = MyWindow()
    constr = MyData()
    window.show()
    sys.exit(app.exec_())
